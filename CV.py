from docx import Document

# Charger le document Word existant
file_path = "C:/Users/a.elouerghi/Desktop/CVV/CV_Achraf ELOUERGHI.docx"
doc = Document(file_path)

# Chercher la section "Distinctions" et ajouter la reconnaissance IEEE
for para in doc.paragraphs:
    if "Distinctions" in para.text:
        index = doc.paragraphs.index(para) + 1
        doc.paragraphs[index].insert_paragraph_before("➡️ Reconnaissance en tant que \"Outstanding Reviewer 2024\" par l'IEEE Instrumentation and Measurement Society.")

# Enregistrer le document modifié
updated_file_path = "C:/Users/a.elouerghi/Desktop/CVV/CV_Achraf ELOUERGHI.docx"
doc.save(updated_file_path)

# Retourner le chemin du fichier mis à jour
updated_file_path
